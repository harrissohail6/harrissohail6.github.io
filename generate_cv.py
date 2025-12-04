from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set margins (ATS-friendly: 0.5-1 inch margins)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Helper function to set font
def set_font(run, name='Calibri', size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    # Set font for Asian text compatibility
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

# Helper function to add horizontal line
def add_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    # Create border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============ HEADER - NAME ============
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('HARIS SOHAIL')
set_font(name_run, 'Calibri', 22, bold=True)
name.paragraph_format.space_after = Pt(2)

# ============ HEADER - TITLE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Electrical Design Engineer')
set_font(title_run, 'Calibri', 12, bold=False)
title.paragraph_format.space_after = Pt(6)

# ============ CONTACT INFO ============
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_text = '+92 312 5331795  |  harrissohail6@gmail.com  |  Rawalpindi, Pakistan  |  linkedin.com/in/haris-sohail-paracha'
contact_run = contact.add_run(contact_text)
set_font(contact_run, 'Calibri', 10)
contact.paragraph_format.space_after = Pt(8)

add_line(doc)

# ============ PROFESSIONAL SUMMARY ============
summary_heading = doc.add_paragraph()
summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
set_font(summary_run, 'Calibri', 12, bold=True)
summary_heading.paragraph_format.space_before = Pt(6)
summary_heading.paragraph_format.space_after = Pt(4)

summary = doc.add_paragraph()
summary_text = 'Detail-oriented Electrical Design Engineer with hands-on experience in electrical drafting, schematic design, and CAD documentation. Currently providing electrical design support for Australian clients through an offshore engineering partnership. Proficient in AutoCAD, MicroStation, and technical documentation with strong expertise in Single Line Diagrams (SLDs), panel layouts, and compliance with Australian electrical standards. Proven ability to coordinate with cross-functional teams and deliver accurate technical drawings within project timelines.'
summary_run = summary.add_run(summary_text)
set_font(summary_run, 'Calibri', 11)
summary.paragraph_format.space_after = Pt(8)

add_line(doc)

# ============ WORK EXPERIENCE ============
exp_heading = doc.add_paragraph()
exp_run = exp_heading.add_run('WORK EXPERIENCE')
set_font(exp_run, 'Calibri', 12, bold=True)
exp_heading.paragraph_format.space_before = Pt(6)
exp_heading.paragraph_format.space_after = Pt(8)

# Job 1 - Advanced Mechanix
job1_title = doc.add_paragraph()
job1_title_run = job1_title.add_run('Electrical Design Engineer')
set_font(job1_title_run, 'Calibri', 11, bold=True)
job1_title.paragraph_format.space_after = Pt(0)

job1_company = doc.add_paragraph()
job1_company_run = job1_company.add_run('Advanced Mechanix (Pvt. Ltd) | Offshore Partner: Overflow Industrial (OFI), Australia')
set_font(job1_company_run, 'Calibri', 11)
job1_company.paragraph_format.space_after = Pt(0)

job1_date = doc.add_paragraph()
job1_date_run = job1_date.add_run('September 2025 - Present')
set_font(job1_date_run, 'Calibri', 10)
job1_date_run.italic = True
job1_date.paragraph_format.space_after = Pt(4)

# Job 1 bullets
bullets1 = [
    'Provide electrical drafting and design support as part of an offshore engineering team partnered with Overflow Industrial (OFI), an Australian-based company',
    'Prepare detailed electrical drawings including schematics, Single Line Diagrams (SLDs), General Arrangement drawings, Block diagrams, and panel layouts using AutoCAD and MicroStation',
    'Coordinate daily with OFI engineers and project managers to understand design requirements and deliver accurate technical documentation',
    'Ensure all drawings comply with Australian electrical standards and client specifications',
    'Manage projects for major Australian clients including CBH Group, Fortescue, and Synergy',
    'Review and incorporate redline markups and as-built modifications into final drawings'
]

for bullet in bullets1:
    p = doc.add_paragraph(style='List Bullet')
    bullet_run = p.add_run(bullet)
    set_font(bullet_run, 'Calibri', 10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

# Job 2 - Professional Hearing Solutions
job2_space = doc.add_paragraph()
job2_space.paragraph_format.space_after = Pt(6)

job2_title = doc.add_paragraph()
job2_title_run = job2_title.add_run('Hearing Instrument Specialist')
set_font(job2_title_run, 'Calibri', 11, bold=True)
job2_title.paragraph_format.space_after = Pt(0)

job2_company = doc.add_paragraph()
job2_company_run = job2_company.add_run('Professional Hearing Solutions (Pvt. Ltd)')
set_font(job2_company_run, 'Calibri', 11)
job2_company.paragraph_format.space_after = Pt(0)

job2_date = doc.add_paragraph()
job2_date_run = job2_date.add_run('September 2024 - September 2025')
set_font(job2_date_run, 'Calibri', 10)
job2_date_run.italic = True
job2_date.paragraph_format.space_after = Pt(4)

bullets2 = [
    'Assisted in hearing assessments, recommended suitable hearing equipment, and provided personalized client consultations',
    'Programmed, fine-tuned, and troubleshot digital hearing aids for optimal performance',
    'Educated clients on device usage, maintenance, and long-term hearing care best practices',
    'Managed client records, follow-ups, and contributed to improving overall service quality'
]

for bullet in bullets2:
    p = doc.add_paragraph(style='List Bullet')
    bullet_run = p.add_run(bullet)
    set_font(bullet_run, 'Calibri', 10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

# Job 3 - HATO Pakistan Internship
job3_space = doc.add_paragraph()
job3_space.paragraph_format.space_after = Pt(6)

job3_title = doc.add_paragraph()
job3_title_run = job3_title.add_run('Production Line Engineer (Internship)')
set_font(job3_title_run, 'Calibri', 11, bold=True)
job3_title.paragraph_format.space_after = Pt(0)

job3_company = doc.add_paragraph()
job3_company_run = job3_company.add_run('HATO Pakistan (Pvt. Ltd)')
set_font(job3_company_run, 'Calibri', 11)
job3_company.paragraph_format.space_after = Pt(0)

job3_date = doc.add_paragraph()
job3_date_run = job3_date.add_run('2023')
set_font(job3_date_run, 'Calibri', 10)
job3_date_run.italic = True
job3_date.paragraph_format.space_after = Pt(4)

bullets3 = [
    'Assisted in monitoring and maintaining production line operations',
    'Supported engineers in diagnosing issues and optimizing machine performance',
    'Applied safety procedures, quality standards, and process documentation practices',
    'Collaborated with technicians on assembly processes and continuous improvement initiatives'
]

for bullet in bullets3:
    p = doc.add_paragraph(style='List Bullet')
    bullet_run = p.add_run(bullet)
    set_font(bullet_run, 'Calibri', 10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

add_line(doc)

# ============ EDUCATION ============
edu_heading = doc.add_paragraph()
edu_run = edu_heading.add_run('EDUCATION')
set_font(edu_run, 'Calibri', 12, bold=True)
edu_heading.paragraph_format.space_before = Pt(6)
edu_heading.paragraph_format.space_after = Pt(8)

# Degree 1
edu1_title = doc.add_paragraph()
edu1_title_run = edu1_title.add_run('Bachelor of Electrical Engineering Technology')
set_font(edu1_title_run, 'Calibri', 11, bold=True)
edu1_title.paragraph_format.space_after = Pt(0)

edu1_school = doc.add_paragraph()
edu1_school_run = edu1_school.add_run('National Skills University Islamabad | 2021 - 2025 | CGPA: 3.89/4.00')
set_font(edu1_school_run, 'Calibri', 10)
edu1_school.paragraph_format.space_after = Pt(6)

# Degree 2
edu2_title = doc.add_paragraph()
edu2_title_run = edu2_title.add_run('F.Sc Pre-Engineering')
set_font(edu2_title_run, 'Calibri', 11, bold=True)
edu2_title.paragraph_format.space_after = Pt(0)

edu2_school = doc.add_paragraph()
edu2_school_run = edu2_school.add_run('Govt. College Asghar Mall, Rawalpindi | 2019 | Grade: B')
set_font(edu2_school_run, 'Calibri', 10)
edu2_school.paragraph_format.space_after = Pt(6)

# Degree 3
edu3_title = doc.add_paragraph()
edu3_title_run = edu3_title.add_run('Matriculation')
set_font(edu3_title_run, 'Calibri', 11, bold=True)
edu3_title.paragraph_format.space_after = Pt(0)

edu3_school = doc.add_paragraph()
edu3_school_run = edu3_school.add_run('Federal Kintaar College Rawalpindi | 2016 | Grade: B')
set_font(edu3_school_run, 'Calibri', 10)
edu3_school.paragraph_format.space_after = Pt(6)

add_line(doc)

# ============ TECHNICAL SKILLS ============
skills_heading = doc.add_paragraph()
skills_run = skills_heading.add_run('TECHNICAL SKILLS')
set_font(skills_run, 'Calibri', 12, bold=True)
skills_heading.paragraph_format.space_before = Pt(6)
skills_heading.paragraph_format.space_after = Pt(8)

skills_data = [
    ('CAD Software:', 'AutoCAD, MicroStation, SkyCAD, Proteus'),
    ('Design Expertise:', 'Single Line Diagrams (SLDs), Schematic Diagrams, Panel Layouts, Block Diagrams, General Arrangement Drawings, Termination Drawings'),
    ('Documentation:', 'Technical Documentation, Redline Markups, As-built Drawings, Project Coordination'),
    ('Standards:', 'Australian Electrical Standards Compliance'),
    ('Soft Skills:', 'Cross-functional Collaboration, Client Communication, Problem Solving, Attention to Detail')
]

for label, value in skills_data:
    p = doc.add_paragraph()
    label_run = p.add_run(label + ' ')
    set_font(label_run, 'Calibri', 10, bold=True)
    value_run = p.add_run(value)
    set_font(value_run, 'Calibri', 10)
    p.paragraph_format.space_after = Pt(3)

add_line(doc)

# ============ PROJECTS ============
proj_heading = doc.add_paragraph()
proj_run = proj_heading.add_run('KEY PROJECT')
set_font(proj_run, 'Calibri', 12, bold=True)
proj_heading.paragraph_format.space_before = Pt(6)
proj_heading.paragraph_format.space_after = Pt(8)

proj_title = doc.add_paragraph()
proj_title_run = proj_title.add_run('SunSync Dual-Axis Solar Tracker with Heat Utilization System')
set_font(proj_title_run, 'Calibri', 11, bold=True)
proj_title.paragraph_format.space_after = Pt(0)

proj_info = doc.add_paragraph()
proj_info_run = proj_info.add_run('Final Year Project | Funded by Ngiri Ignite 2023-24')
set_font(proj_info_run, 'Calibri', 10)
proj_info_run.italic = True
proj_info.paragraph_format.space_after = Pt(4)

proj_bullets = [
    'Designed and developed an innovative dual-axis solar tracking system to maximize solar energy capture through intelligent sun-following technology',
    'Implemented a heat utilization system to capture and repurpose excess thermal energy for secondary applications',
    'Applied embedded systems programming and control systems principles for precise motor control and sensor integration',
    'Demonstrated practical application of electrical engineering principles in renewable energy solutions'
]

for bullet in proj_bullets:
    p = doc.add_paragraph(style='List Bullet')
    bullet_run = p.add_run(bullet)
    set_font(bullet_run, 'Calibri', 10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

add_line(doc)

# ============ KEY CLIENTS ============
clients_heading = doc.add_paragraph()
clients_run = clients_heading.add_run('KEY CLIENTS SERVED')
set_font(clients_run, 'Calibri', 12, bold=True)
clients_heading.paragraph_format.space_before = Pt(6)
clients_heading.paragraph_format.space_after = Pt(4)

clients = doc.add_paragraph()
clients_run = clients.add_run('CBH Group  |  Fortescue  |  Synergy  |  Overflow Industrial (OFI)')
set_font(clients_run, 'Calibri', 10)
clients.paragraph_format.space_after = Pt(6)

# Save document
doc.save(r'C:\Users\Lenovo\HarisPortfolio\Haris_Sohail_CV.docx')
print('CV created successfully!')
